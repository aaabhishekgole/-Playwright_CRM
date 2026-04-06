"""
Generate Word (.docx) portal guides for Admin and Pickup Runner.
Run: python generate_docs.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "screenshots")

# ── colour palette (stored as (R, G, B) tuples) ───────────────────────────────
BRAND_DARK   = (0x07, 0x0D, 0x1A)
BRAND_BLUE   = (0x00, 0x78, 0xD4)
BRAND_CYAN   = (0x00, 0xBC, 0xD4)
BRAND_GREEN  = (0x00, 0xC8, 0x53)
BRAND_ORANGE = (0xFF, 0x67, 0x00)
WHITE        = (0xFF, 0xFF, 0xFF)
LIGHT_GREY   = (0xF0, 0xF4, 0xF8)
MID_GREY     = (0xCC, 0xD6, 0xE0)
TEXT_DARK    = (0x1A, 0x20, 0x2C)
TEXT_MUTED   = (0x55, 0x65, 0x78)

def rgb(t): return RGBColor(t[0], t[1], t[2])
def hx(t):  return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# ── low-level XML helpers ─────────────────────────────────────────────────────

def set_cell_bg(cell, color_tuple):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hx(color_tuple))
    tcPr.append(shd)


def set_cell_border(cell, **sides):
    tcPr    = cell._tc.get_or_add_tcPr()
    tcBdrs  = OxmlElement('w:tcBorders')
    for side, cfg in sides.items():
        if cfg:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   cfg.get('val',   'single'))
            el.set(qn('w:sz'),    str(cfg.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), cfg.get('color', '000000'))
            tcBdrs.append(el)
    tcPr.append(tcBdrs)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBdrs = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBdrs.append(el)
    tblPr.append(tblBdrs)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def set_doc_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0):
    sec = doc.sections[0]
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bottom)
    sec.left_margin   = Cm(left)
    sec.right_margin  = Cm(right)


# ── paragraph / run helpers ───────────────────────────────────────────────────

def add_left_border_para(doc, fill_hex, border_color_hex, indent_cm=0.3):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)
    pBdr  = OxmlElement('w:pBdr')
    left  = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), border_color_hex)
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent   = Cm(indent_cm)
    p.paragraph_format.space_before  = Pt(14)
    p.paragraph_format.space_after   = Pt(6)
    return p


def add_page_divider(doc, color=BRAND_BLUE):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), hx(color))
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)


# ── high-level components ─────────────────────────────────────────────────────

def add_section_heading(doc, number, text, color=BRAND_BLUE):
    p = add_left_border_para(doc,
        fill_hex=hx(LIGHT_GREY),
        border_color_hex=hx(color))
    num_r = p.add_run(f"{number}.  ")
    num_r.font.size  = Pt(14)
    num_r.font.bold  = True
    num_r.font.color.rgb = rgb(color)
    txt_r = p.add_run(text)
    txt_r.font.size  = Pt(14)
    txt_r.font.bold  = True
    txt_r.font.color.rgb = rgb(TEXT_DARK)


def add_sub_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size  = Pt(12 if level == 2 else 11)
    r.font.bold  = True
    r.font.color.rgb = rgb(BRAND_DARK if level == 2 else TEXT_MUTED)


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size  = Pt(10.5)
    r.font.color.rgb = rgb(TEXT_DARK)
    r.italic = italic
    return p


def add_note(doc, text):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'E8F9FB')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), hx(BRAND_CYAN))
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    label = p.add_run("  Note:  ")
    label.font.bold  = True
    label.font.size  = Pt(10)
    label.font.color.rgb = rgb(BRAND_CYAN)
    body = p.add_run(text)
    body.font.size   = Pt(10)
    body.font.italic = True
    body.font.color.rgb = rgb(TEXT_DARK)


def add_styled_table(doc, headers, rows, col_widths=None, header_color=BRAND_BLUE):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, header_color)
        set_cell_border(c, bottom={'val': 'single', 'sz': 4, 'color': 'FFFFFF'})
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(10)
        r.font.color.rgb = rgb(WHITE)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    sep_color = hx(MID_GREY)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = LIGHT_GREY if ri % 2 == 0 else WHITE
        for ci, text in enumerate(row_data):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            set_cell_border(c, bottom={'val': 'single', 'sz': 2, 'color': sep_color})
            p = c.paragraphs[0]
            r = p.add_run(text)
            r.font.size  = Pt(10)
            r.font.bold  = (ci == 0)
            r.font.color.rgb = rgb(TEXT_DARK)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        set_col_widths(table, col_widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_checklist_table(doc, items, check_color=BRAND_GREEN):
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    set_col_widths(table, [1.0, 13.0])
    for i, item in enumerate(items):
        bg = LIGHT_GREY if i % 2 == 0 else WHITE
        cb, tb = table.rows[i].cells[0], table.rows[i].cells[1]
        set_cell_bg(cb, bg); set_cell_bg(tb, bg)
        cp = cb.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run("☐")
        cr.font.size  = Pt(12)
        cr.font.color.rgb = rgb(check_color)
        tr = tb.paragraphs[0].add_run(item)
        tr.font.size  = Pt(10.5)
        tr.font.color.rgb = rgb(TEXT_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_screenshot(doc, filename, caption="", width_inches=5.8):
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        add_body(doc, f"[Screenshot not found: {filename}]", italic=True)
        return
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    border_cfg = {'val': 'single', 'sz': 4, 'color': hx(BRAND_BLUE)}
    set_cell_border(cell, top=border_cfg, bottom=border_cfg,
                    left=border_cfg, right=border_cfg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.size   = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = rgb(TEXT_MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_cover_page(doc, title, subtitle, role_label, role_color):
    # brand bar
    for text, size, color, bold in [
        ("GADGET SEVA HUB", 13, BRAND_CYAN,  True),
        ("Enterprise Service Operations", 10, TEXT_MUTED, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = rgb(color)
        if bold:
            r.font.all_caps = True

    doc.add_paragraph()

    # role badge via 1-cell table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    cell.width = Cm(6)
    set_cell_bg(cell, role_color)
    rp = cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(f"  {role_label}  ")
    rr.font.size  = Pt(11)
    rr.font.bold  = True
    rr.font.color.rgb = rgb(WHITE)

    doc.add_paragraph()

    # title
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(title)
    rt.font.size  = Pt(28)
    rt.font.bold  = True
    rt.font.color.rgb = rgb(BRAND_DARK)

    # subtitle
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run(subtitle)
    rs.font.size  = Pt(12)
    rs.font.color.rgb = rgb(TEXT_MUTED)

    doc.add_paragraph()

    # divider
    p_div = doc.add_paragraph()
    pPr   = p_div._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), hx(BRAND_BLUE))
    pBdr.append(bot)
    pPr.append(pBdr)

    # version
    pv = doc.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = pv.add_run("Version: April 2026  |  Gadget Seva Hub Ops Console")
    rv.font.size  = Pt(9)
    rv.font.color.rgb = rgb(TEXT_MUTED)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# ADMIN GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def build_admin_guide():
    doc = Document()
    set_doc_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    add_cover_page(doc,
        title="Admin Portal User Guide",
        subtitle="Complete step-by-step guide for administrators\nto operate the Gadget Seva Hub Ops Console",
        role_label="ROLE: ADMIN",
        role_color=BRAND_BLUE)

    # 1 — Login
    add_section_heading(doc, 1, "How to Log In")
    add_body(doc, "Open the portal URL in your browser. You will see the Sign in to Ops Console screen.")
    add_screenshot(doc, "01-login-page.png", "Figure 1 — Login Screen")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Open browser and navigate to the portal URL"],
         ["2","Enter your Username (e.g. admin)"],
         ["3","Enter your Password (e.g. Admin@123)"],
         ["4","Click the Enter Console button"]],
        col_widths=[1.8, 12.2])
    add_note(doc, "The login page shows all seeded usernames and the default password Admin@123 in the LOCAL ACCESS tile.")

    # 2 — Dashboard
    add_page_divider(doc)
    add_section_heading(doc, 2, "Dashboard — Your Home Screen")
    add_body(doc, "After logging in you land on the Dashboard — your operational overview showing live KPIs across all workflow stages.")
    add_screenshot(doc, "02-admin-dashboard.png", "Figure 2 — Admin Dashboard")
    add_styled_table(doc,
        ["Section", "Purpose"],
        [["KPI Cards (top row)","Total requests, open, in-progress, and completed counts at a glance"],
         ["SLA Status","How many requests are within SLA vs breached"],
         ["Recent Activity","Latest status changes across all claims"],
         ["Left Sidebar","Full navigation — expand any section to go to that module"]],
        col_widths=[4.5, 9.5])
    add_note(doc, "Click any section in the left sidebar to expand it, then click a sub-item to navigate.")

    # 3 — Create Request
    add_page_divider(doc)
    add_section_heading(doc, 3, "Creating a Service Request")
    add_body(doc, "Register a new customer claim. Go to: Service Requests → Create Request")
    add_screenshot(doc, "03-admin-create-request.png", "Figure 3 — Create Service Request Form")
    add_styled_table(doc,
        ["Step", "Field", "What to Enter"],
        [["1","Customer Name","Full name of the customer or company"],
         ["2","Phone","Primary contact number"],
         ["3","Address","Pickup address — Line 1, City, State, Postal Code"],
         ["4","Device Brand / Model","e.g. Samsung Galaxy S21"],
         ["5","Device Category","Select from dropdown (Mobile, Laptop, Tablet, etc.)"],
         ["6","Serial Number","Device serial or IMEI number"],
         ["7","Issue Summary","Brief one-line description of the problem"],
         ["8","Priority","Select LOW / MEDIUM / HIGH / CRITICAL"],
         ["9","Source Channel","How the request came in (Walk-in, WhatsApp, Email, etc.)"],
         ["10","Submit","Click Submit — ticket created with status: REGISTERED"]],
        col_widths=[1.5, 4.0, 8.5])

    # 4 — All Requests
    add_page_divider(doc)
    add_section_heading(doc, 4, "Managing All Requests")
    add_body(doc, "Go to: Service Requests → All Requests. Shows every claim across all statuses.")
    add_screenshot(doc, "04-admin-all-requests.png", "Figure 4 — All Requests Queue")
    add_styled_table(doc,
        ["Action", "How"],
        [["Search a claim","Use the search box — by ticket number, customer name, or device"],
         ["Filter by status","Use the status filter dropdown"],
         ["Open a claim","Click any request card to open the full details page"],
         ["View timeline","Inside a request, scroll to the Timeline tab"],
         ["Upload documents","Inside a request, use the Attachments tab"]],
        col_widths=[5.0, 9.0])
    add_sub_heading(doc, "Request Status Flow", level=2)
    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.left_indent = Cm(0.5)
    for s in ["REGISTERED","→ PICKUP_ASSIGNED","→ PICKED_UP","→ RECEIVED_AT_HUB",
              "→ DIAGNOSIS_IN_PROGRESS","→ ESTIMATE_SUBMITTED","→ ESTIMATE_APPROVED",
              "→ REPAIR_IN_PROGRESS","→ REPAIR_COMPLETED","→ QC_PASSED",
              "→ READY_FOR_DISPATCH","→ OUT_FOR_DELIVERY","→ DELIVERED","→ INVOICED","→ CLOSED"]:
        r = p_flow.add_run(s + "  ")
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb(BRAND_BLUE if "→" not in s else TEXT_DARK)

    # 5 — Pickup Management
    add_page_divider(doc)
    add_section_heading(doc, 5, "Pickup Management")

    add_sub_heading(doc, "5.1  Pickup Dashboard")
    add_body(doc, "Go to: Pickup Management → Pickup Dashboard. Live count of pickups in each stage.")
    add_screenshot(doc, "05-admin-pickup-dashboard.png", "Figure 5 — Pickup Dashboard")

    add_sub_heading(doc, "5.2  Runner Onboarding")
    add_body(doc, "Add new pickup runners before assigning them. Go to: Pickup Management → Runner Onboarding")
    add_screenshot(doc, "06-admin-runner-onboarding.png", "Figure 6 — Runner Onboarding Form")
    add_styled_table(doc,
        ["Step", "Field", "Detail"],
        [["1","Full Name","Runner's legal full name"],
         ["2","Phone","Mobile number used for WhatsApp link delivery"],
         ["3","WhatsApp Number","If different from phone number"],
         ["4","Email","Optional contact email"],
         ["5","Username","Login username for the runner app"],
         ["6","Save Runner","Click to create the account"]],
        col_widths=[1.5, 4.0, 8.5])

    add_sub_heading(doc, "5.3  Assign Pickup")
    add_body(doc, "Assign a runner to a registered request. Go to: Pickup Management → Assign Pickup")
    add_screenshot(doc, "07-admin-assign-pickup.png", "Figure 7 — Assign Pickup")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Find the request in the list (use search or scroll)"],
         ["2","Click Assign Pickup on the request card"],
         ["3","Select the Runner from the dropdown"],
         ["4","Set the Scheduled Date and Time for the pickup"],
         ["5","Add any Notes for the runner (optional)"],
         ["6","Click Save Assignment"]],
        col_widths=[1.8, 12.2])
    add_note(doc, "The runner receives a WhatsApp link to the pickup portal automatically after assignment.")

    add_sub_heading(doc, "5.4  Pending Pickup")
    add_body(doc, "View all pickups assigned but not yet completed. Go to: Pickup Management → Pending Pickup")
    add_screenshot(doc, "08-admin-pending-pickup.png", "Figure 8 — Pending Pickup Queue")

    # 6 — Users
    add_page_divider(doc)
    add_section_heading(doc, 6, "User Management")
    add_body(doc, "Manage all portal users. Go to: Users section in the left sidebar.")
    add_screenshot(doc, "09-admin-users.png", "Figure 9 — User Management")
    add_styled_table(doc,
        ["Menu Item", "Purpose"],
        [["Admin Users","Create and manage admin accounts"],
         ["Delivery Executives","Manage delivery agent accounts"],
         ["Hub Operators","Manage hub inward staff accounts"],
         ["Service Center Users","Manage technician accounts"],
         ["Customers","View customer directory and profiles"],
         ["Roles & Permissions","View the access control matrix per role"]],
        col_widths=[5.0, 9.0])
    add_note(doc, "To create a user: click the user type → Add User → fill Name, Phone, Username, Role → Save.")

    # 7 — Billing
    add_page_divider(doc)
    add_section_heading(doc, 7, "Billing & Invoicing")
    add_body(doc, "Generate GST-compliant invoices and track payments. Go to: Billing section.")
    add_screenshot(doc, "10-admin-billing.png", "Figure 10 — Billing & Invoicing")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Open the request that has been delivered and repair is complete"],
         ["2","Go to Billing → Generate Invoice"],
         ["3","Select the request from the list"],
         ["4","Fill in GST details (State Code, Place of Supply, GST Rate)"],
         ["5","Add Labour and Parts descriptions"],
         ["6","Click Generate Invoice — PDF is created and linked to the request"]],
        col_widths=[1.8, 12.2])
    add_note(doc, "For Payment Reconciliation: Billing → Payment Reconciliation → find payment by UTR → mark RECONCILED or MISMATCHED → save.")

    # 8 — Documents
    add_page_divider(doc)
    add_section_heading(doc, 8, "Document Library")
    add_body(doc, "Upload and share SOPs, policies, training materials, and reports. Go to: Documents → Document Library")
    add_screenshot(doc, "11-admin-documents.png", "Figure 11 — Document Library")
    add_styled_table(doc,
        ["Step", "Field", "Detail"],
        [["1","Name","Document title (e.g. Pickup SOP v2)"],
         ["2","Category","Select SOP / Policy / Training / Template / Report / Other"],
         ["3","Description","Brief one-line description"],
         ["4","File","Click to choose a file (PDF, image, Word, Excel, etc.)"],
         ["5","Upload Document","Click to store the file and generate a secure link"]],
        col_widths=[1.5, 4.0, 8.5])
    add_styled_table(doc,
        ["Action", "How"],
        [["View","Click View — opens PDF or image in a new browser tab"],
         ["Download","Click Download — saves the file to your computer"],
         ["Delete","Click Delete (Admin only) → confirm the prompt"],
         ["Filter by category","Use the category dropdown at the top of the list"]],
        col_widths=[4.0, 10.0],
        header_color=BRAND_CYAN)

    # 9 — Reports
    add_page_divider(doc)
    add_section_heading(doc, 9, "Reports")
    add_body(doc, "View operational reports. Go to: Reports section in the left sidebar.")
    add_screenshot(doc, "12-admin-reports.png", "Figure 12 — Reports")
    add_styled_table(doc,
        ["Report", "What It Shows"],
        [["Request Report","Intake volume, open vs closed counts by period"],
         ["Pickup Report","Pickup completion rates and runner performance"],
         ["Repair Report","Repair throughput and total loss rates"],
         ["Delivery Report","Dispatch and completion metrics"],
         ["SLA / TAT Report","Breach counts and average turnaround time"],
         ["Revenue Report","Collections, outstanding amounts, refunds"],
         ["Audit Logs","Full enterprise change log"]],
        col_widths=[5.0, 9.0])

    # 10 — Notifications
    add_page_divider(doc)
    add_section_heading(doc, 10, "Notifications")
    add_body(doc, "Monitor SMS and email delivery to customers. Go to: Notifications section.")
    add_screenshot(doc, "13-admin-notifications.png", "Figure 13 — Notifications")
    add_styled_table(doc,
        ["Section", "Purpose"],
        [["SMS Logs","All outbound SMS messages with delivery status"],
         ["Email Logs","All outbound emails with delivery status"],
         ["Failed Notifications","Messages that failed — can be retried"],
         ["Templates","View and manage notification message templates"]],
        col_widths=[5.0, 9.0])
    add_note(doc, "To retry: Notifications → Failed Notifications → find the message → click Retry.")

    # 11 — Settings
    add_page_divider(doc)
    add_section_heading(doc, 11, "Settings")
    add_body(doc, "Configure system-wide behaviour. Go to: Settings section.")
    add_screenshot(doc, "14-admin-settings.png", "Figure 14 — Settings")
    add_styled_table(doc,
        ["Section", "Purpose"],
        [["Status Configuration","Manage workflow states and allowed transitions"],
         ["Notification Settings","Configure retry intervals and channel priority"],
         ["SLA Configuration","Set SLA hours per tenant and priority level"],
         ["File Storage Config","Configure storage paths and signed URL TTL"],
         ["System Preferences","Global platform preferences"]],
        col_widths=[5.0, 9.0])
    add_note(doc, "Changes to status configuration affect the entire workflow. Make changes only during non-peak hours.")

    # 12 — Audit
    add_page_divider(doc)
    add_section_heading(doc, 12, "Audit Logs")
    add_body(doc, "Full traceability of every action taken in the system. Go to: Audit section.")
    add_screenshot(doc, "15-admin-audit.png", "Figure 15 — Audit Logs")
    add_styled_table(doc,
        ["Section", "What It Tracks"],
        [["Activity Logs","All actions across every module (create, update, delete)"],
         ["Status History","Every status transition for every request"],
         ["User Actions","Actions performed by each individual user"],
         ["Change Logs","Before and after values for every field change"]],
        col_widths=[5.0, 9.0])

    # 13 — Quick Reference
    add_page_divider(doc)
    add_section_heading(doc, 13, "Quick Reference")
    add_styled_table(doc,
        ["Task", "Where to Go"],
        [["Create a new claim","Service Requests → Create Request"],
         ["Find a specific ticket","Service Requests → Search Request"],
         ["Assign a pickup","Pickup Management → Assign Pickup"],
         ["Add a new runner","Pickup Management → Runner Onboarding"],
         ["Generate an invoice","Billing → Generate Invoice"],
         ["Reconcile a payment","Billing → Payment Reconciliation"],
         ["Add a new staff member","Users → (select role)"],
         ["Upload an SOP or policy","Documents → Document Library"],
         ["Check SLA breaches","Dashboard → SLA / TAT Summary"],
         ["View audit trail","Audit → Activity Logs"],
         ["Retry a failed SMS","Notifications → Failed Notifications"]],
        col_widths=[6.0, 8.0],
        header_color=BRAND_DARK)

    out = os.path.join(BASE, "Admin-Portal-User-Guide.docx")
    doc.save(out)
    print(f"  Saved: {out}")


# ═══════════════════════════════════════════════════════════════════════════════
# PICKUP RUNNER GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def build_pickup_guide():
    doc = Document()
    set_doc_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    add_cover_page(doc,
        title="Pickup Runner Portal Guide",
        subtitle="Step-by-step field guide for pickup runners\nto execute pickups using the Gadget Seva Hub portal",
        role_label="ROLE: PICKUP_AGENT",
        role_color=BRAND_ORANGE)

    # 1 — Login
    add_section_heading(doc, 1, "How to Log In", color=BRAND_ORANGE)
    add_body(doc, "Open the portal in your browser and enter your credentials.")
    add_screenshot(doc, "01-login-page.png", "Figure 1 — Login Screen")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Open browser and navigate to the portal URL"],
         ["2","Enter your Username (given by Admin — e.g. pickup)"],
         ["3","Enter your Password (default: Admin@123)"],
         ["4","Click Enter Console"]],
        col_widths=[1.8, 12.2],
        header_color=BRAND_ORANGE)
    add_note(doc, "If you cannot log in, contact your Admin to verify your account is active.")

    # 2 — Dashboard
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 2, "Your Dashboard After Login", color=BRAND_ORANGE)
    add_body(doc, "After logging in you see the Dashboard. Your main working area is the Pickup Management section in the left sidebar.")
    add_screenshot(doc, "16-pickup-dashboard.png", "Figure 2 — Pickup Runner Dashboard")
    add_styled_table(doc,
        ["Left Menu Section", "What You Use It For"],
        [["Dashboard","Overview and KPI summary"],
         ["Service Requests","View all requests relevant to your role"],
         ["Pickup Management","Your active daily workflow — pending, picked up, failed, history"]],
        col_widths=[5.0, 9.0],
        header_color=BRAND_ORANGE)
    add_note(doc, "You will only see sections relevant to PICKUP_AGENT role. Billing, Settings, and Admin screens are not accessible.")

    # 3 — Pending Pickup
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 3, "Pending Pickup — Your Active Assignments", color=BRAND_ORANGE)
    add_body(doc, "Go to: Pickup Management → Pending Pickup. Shows every pickup assigned to you that is not yet complete.")
    add_screenshot(doc, "17-pickup-pending.png", "Figure 3 — Pending Pickup Queue")
    add_styled_table(doc,
        ["Field", "Meaning"],
        [["Request Number","Ticket reference (e.g. GSH-2026-0001)"],
         ["Customer Name","Name of the customer you are collecting from"],
         ["Address","Pickup location"],
         ["Scheduled Time","When you are expected to reach the customer"],
         ["Status","PICKUP_ASSIGNED or PICKUP_ACCEPTED"],
         ["OTP","4-digit code — verify with customer BEFORE collecting the device"]],
        col_widths=[4.0, 10.0],
        header_color=BRAND_ORANGE)

    # 4 — Picked Up Devices
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 4, "Picked Up Devices — Devices You Have Collected", color=BRAND_ORANGE)
    add_body(doc, "Go to: Pickup Management → Picked Up Devices. Shows all devices you have collected and photographed.")
    add_screenshot(doc, "18-pickup-picked-up.png", "Figure 4 — Picked Up Devices")
    add_note(doc, "Once a device shows here, the hub team takes over. No further action needed from you.")

    # 5 — Failed Cases
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 5, "Pickup Failed Cases — Failed Attempts", color=BRAND_ORANGE)
    add_body(doc, "Go to: Pickup Management → Pickup Failed Cases. If a pickup attempt fails, it appears here.")
    add_screenshot(doc, "19-pickup-failed.png", "Figure 5 — Pickup Failed Cases")
    add_styled_table(doc,
        ["Reason", "What to Do"],
        [["Customer not home","Contact Admin — they will reschedule"],
         ["Customer refused pickup","Contact Admin — request needs escalation"],
         ["Wrong address","Contact Admin with the correct address details"],
         ["Device not ready","Contact Admin — they will update the schedule"]],
        col_widths=[5.0, 9.0],
        header_color=BRAND_ORANGE)

    # 6 — History
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 6, "Pickup History — Your Completed Pickups", color=BRAND_ORANGE)
    add_body(doc, "Go to: Pickup Management → Pickup History. Your full record of every successfully completed pickup.")
    add_screenshot(doc, "20-pickup-history.png", "Figure 6 — Pickup History")
    add_body(doc, "Use this screen to:\n• Verify a pickup was recorded correctly\n• Check the date and time of a past collection\n• Look up which photos were uploaded for a specific device")

    # 7 — Runner App
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 7, "Runner App — Field Pickup Flow (Step by Step)", color=BRAND_ORANGE)
    add_body(doc, "The Runner App is your main tool at the customer's location. Works on mobile and desktop browsers.")
    add_screenshot(doc, "21-runner-app.png", "Figure 7 — Runner App")

    add_sub_heading(doc, "Step 1 — Log in to the Runner App")
    add_body(doc, "Open: https://front-end-uat.up.railway.app/runner-app  on your phone browser. Enter your username and password.")

    add_sub_heading(doc, "Step 2 — Accept the Pickup Assignment")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Arrive at the customer's location"],
         ["2","Open the Runner App on your phone"],
         ["3","Find the correct assignment in your list"],
         ["4","Tap Accept Pickup — status changes to PICKUP_ACCEPTED"],
         ["5","Verify the 4-digit OTP with the customer before touching the device"]],
        col_widths=[1.8, 12.2],
        header_color=BRAND_ORANGE)
    add_note(doc, "Ask the customer for the OTP from their SMS/WhatsApp. Match it with the OTP in your portal before proceeding.")

    add_sub_heading(doc, "Step 3 — Upload 10 Mandatory Device Photos")
    add_body(doc, "You must upload ALL 10 photos before you can complete the pickup.")
    add_styled_table(doc,
        ["#", "Photo Type", "What to Capture"],
        [["1","Front","Screen facing forward, full device visible"],
         ["2","Back","Back panel, full device visible"],
         ["3","Left Side","Left edge of the device"],
         ["4","Right Side","Right edge of the device"],
         ["5","Top","Top edge (power button, headphone jack area)"],
         ["6","Bottom","Bottom edge (charging port area)"],
         ["7","Display ON","Screen powered on showing the home screen"],
         ["8","Serial Label","IMEI sticker or serial number label"],
         ["9","Damage Closeup","Any existing visible damage zoomed in"],
         ["10","Accessories","All accessories handed over (charger, case, cables)"]],
        col_widths=[1.2, 4.0, 8.8],
        header_color=BRAND_ORANGE)
    add_note(doc, "Do NOT skip any photo. The system will NOT allow you to mark the pickup as complete if any photo is missing.")

    add_sub_heading(doc, "Step 4 — Complete the Pickup")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Verify all 10 photos show green ticks in the portal"],
         ["2","Scroll to the bottom of the pickup portal page"],
         ["3","Tap Complete Pickup"],
         ["4","Confirm the action in the prompt"],
         ["5","Status changes to PICKED_UP — you are done"]],
        col_widths=[1.8, 12.2],
        header_color=BRAND_ORANGE)

    # 8 — WhatsApp Link
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 8, "Using the Pickup Portal Link (WhatsApp)", color=BRAND_ORANGE)
    add_body(doc, "When Admin assigns a pickup, you receive a WhatsApp message with a unique portal link.")
    add_styled_table(doc,
        ["Step", "Action"],
        [["1","Open the WhatsApp message from Gadget Seva Hub"],
         ["2","Tap the link — it opens the pickup portal in your phone browser"],
         ["3","No login required — the link is pre-authenticated for you"],
         ["4","Follow Steps 2–4 above (accept → upload photos → complete)"]],
        col_widths=[1.8, 12.2],
        header_color=BRAND_ORANGE)
    add_note(doc, "The link is valid for your assigned pickup only. Do not share it with others.")

    # 9 — Failure Handling
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 9, "What to Do If Pickup Fails", color=BRAND_ORANGE)
    add_styled_table(doc,
        ["Situation", "Action"],
        [["Customer not available","Mark PICKUP_FAILED, add reason in Notes, contact Admin"],
         ["Customer refused","Mark PICKUP_FAILED, note 'Customer refused', contact Admin"],
         ["Address not found","Mark PICKUP_FAILED, note 'Wrong address', share correct location"],
         ["Device not in condition","Mark PICKUP_FAILED, note the condition issue"]],
        col_widths=[5.5, 8.5],
        header_color=BRAND_ORANGE)
    add_sub_heading(doc, "How to mark a pickup as failed:", level=3)
    add_body(doc, "1. Open the pickup portal link or go to Pending Pickup\n2. Tap Update Status\n3. Select PICKUP_FAILED from the dropdown\n4. Add clear notes describing the reason\n5. Tap Save")
    add_note(doc, "Always call the customer before marking as failed. Wait up to 30 minutes if they need time.")

    # 10 — Quick Reference + Checklist
    add_page_divider(doc, color=BRAND_ORANGE)
    add_section_heading(doc, 10, "Quick Reference & Mandatory Photo Checklist", color=BRAND_ORANGE)
    add_styled_table(doc,
        ["Task", "Where"],
        [["See my assigned pickups","Pickup Management → Pending Pickup"],
         ["Accept a pickup on-site","Runner App → Accept Pickup"],
         ["Upload device photos","Runner App → Photos section"],
         ["Complete a pickup","Runner App → Complete Pickup"],
         ["View collected devices","Pickup Management → Picked Up Devices"],
         ["Report a failed pickup","Runner App → Update Status → PICKUP_FAILED"],
         ["View my history","Pickup Management → Pickup History"],
         ["Log in via WhatsApp link","Tap the link in your WhatsApp message"]],
        col_widths=[6.0, 8.0],
        header_color=BRAND_ORANGE)

    add_sub_heading(doc, "Mandatory Photo Checklist (use at every pickup)")
    add_checklist_table(doc, [
        "Front of device",
        "Back of device",
        "Left side",
        "Right side",
        "Top edge",
        "Bottom edge",
        "Screen powered ON",
        "Serial number / IMEI label",
        "Any existing damage (closeup)",
        "All accessories included",
    ], check_color=BRAND_ORANGE)

    add_sub_heading(doc, "Important Rules")
    add_styled_table(doc,
        ["#", "Rule"],
        [["1","Always verify the OTP before collecting any device"],
         ["2","Never skip any of the 10 photos — incomplete evidence delays the claim"],
         ["3","Report failures immediately — do not leave without updating the portal"],
         ["4","Do not share your WhatsApp portal link — it is unique to your assignment"],
         ["5","Contact Admin for any disputes — do not make promises to customers about timelines"]],
        col_widths=[1.2, 12.8],
        header_color=BRAND_DARK)

    out = os.path.join(BASE, "Pickup-Runner-Portal-Guide.docx")
    doc.save(out)
    print(f"  Saved: {out}")


# ── entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating Word documents...")
    build_admin_guide()
    build_pickup_guide()
    print("Done.")
